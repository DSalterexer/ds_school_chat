import streamlit as st
from langchain_openai import OpenAIEmbeddings, OpenAI
from langchain.chains import RetrievalQA
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
import tempfile
import os
from docx import Document

# твой OpenAI API ключ
OPENAI_API_KEY = "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"

st.title("Веб-чат с загрузкой и выгрузкой файлов")

# Загрузка файла
uploaded_file = st.file_uploader("Загрузи документ (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name

    # загрузка документов
    if uploaded_file.type == "application/pdf":
        loader = PyPDFLoader(tmp_path)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        loader = Docx2txtLoader(tmp_path)
    else:
        loader = TextLoader(tmp_path, encoding='utf-8')

    documents = loader.load()
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    texts = text_splitter.split_documents(documents)

    embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    vectordb = Chroma.from_documents(texts, embeddings)

    qa_bot = RetrievalQA.from_chain_type(
        llm=OpenAI(model="gpt-3.5-turbo", temperature=0, api_key=OPENAI_API_KEY),
        chain_type="stuff",
        retriever=vectordb.as_retriever(search_kwargs={"k": 4})
    )

    question = st.text_input("Задай вопрос по документу:")

    if question:
        result = qa_bot({"query": question})
        st.write("Ответ:", result["result"])

        # Создание файла для скачивания
        doc = Document()
        doc.add_heading("Ответ на вопрос:", level=1)
        doc.add_paragraph(question)
        doc.add_heading("Результат:", level=2)
        doc.add_paragraph(result["result"])

        output_path = "result.docx"
        doc.save(output_path)

        # Скачивание файла пользователем
        with open(output_path, "rb") as f:
            st.download_button(
                label="Скачать ответ в виде DOCX-файла",
                data=f,
                file_name="result.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Удаление временных файлов
    os.unlink(tmp_path)
